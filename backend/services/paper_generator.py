"""
Question Paper Generator Module
Generates PDF and DOCX question papers from question banks using blueprints
"""

import json
import re
import tempfile
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any
from uuid import uuid4
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
from core.database import get_db_type
from services.image_integration import (
    detect_image_required_in_question,
    get_image_for_question,
    save_image_blob_to_temp,
    cleanup_temp_image_file,
    fix_extreme_brightness_image,
    rotate_image_for_pdf_insertion
)
from services.rag_config import logger


NUMBER_WORDS = {
    "one": 1,
    "two": 2,
    "three": 3,
    "four": 4,
    "five": 5,
    "six": 6,
    "seven": 7,
    "eight": 8,
    "nine": 9,
    "ten": 10,
}

BLOOMS_ORDER = ["R", "U", "A", "AN", "E", "C"]
BLOOMS_LABEL = {
    "R": "Remember",
    "U": "Understand",
    "A": "Apply",
    "AN": "Analyze",
    "E": "Evaluate",
    "C": "Create",
}


def _blooms_code_from_marks(marks: float) -> str:
    try:
        value = float(marks)
    except Exception:
        value = 0.0

    if value <= 1:
        return "R"
    if value <= 2:
        return "U"
    if value <= 5:
        return "A"
    if value <= 10:
        return "AN"
    if value <= 14:
        return "E"
    return "C"


def _blooms_code_from_level(level: str | None) -> str | None:
    if not level:
        return None

    normalized = str(level).strip().lower()
    if not normalized:
        return None

    level_map = {
        "remember": "R",
        "understand": "U",
        "apply": "A",
        "analyze": "AN",
        "analyse": "AN",
        "evaluate": "E",
        "create": "C",
    }

    # If multiple levels are provided, take the first one for mapping.
    first_level = normalized.split(",")[0].strip()
    return level_map.get(first_level)


def _build_blooms_mapping(blueprint: Dict[str, Any], questions_by_part: Dict[str, List[Dict[str, Any]]]) -> Dict[str, Dict[str, Any]]:
    mapping = {
        code: {"questions": [], "marks": 0.0, "label": BLOOMS_LABEL[code]}
        for code in BLOOMS_ORDER
    }

    question_number = 1
    for part in blueprint.get('parts', []):
        part_name = part.get('part_name') or part.get('name')
        part_questions = questions_by_part.get(part_name, [])
        part_marks = float(part.get('marks_per_question') or 0)

        for q in part_questions:
            q_marks = float(q.get('marks') or part_marks)
            code = _blooms_code_from_level(q.get('blooms_level')) or _blooms_code_from_marks(q_marks)
            mapping[code]["questions"].append(question_number)
            mapping[code]["marks"] += q_marks
            question_number += 1

    return mapping


def _parse_answer_any_count(instruction: str) -> int | None:
    text = (instruction or "").strip().lower()
    if not text:
        return None

    match = re.search(r"answer\s+any\s+(\d+|one|two|three|four|five|six|seven|eight|nine|ten)", text)
    if not match:
        return None

    token = match.group(1)
    if token.isdigit():
        return int(token)
    return NUMBER_WORDS.get(token)


def _effective_answer_count(part: Dict[str, Any], fetched_questions_count: int) -> int:
    configured_count = int(part.get('num_questions') or part.get('count') or fetched_questions_count or 0)
    instruction = part.get('instructions') or part.get('instruction') or ""
    answer_any = _parse_answer_any_count(instruction)

    if answer_any is not None:
        return max(0, min(answer_any, configured_count, fetched_questions_count))
    return max(0, min(configured_count, fetched_questions_count))


def _resolve_course_outcome_image(course_outcome_file: str) -> tuple[str | None, str | None]:
    """
    Returns (image_path_to_render, temp_file_to_cleanup)
    - If already image: returns original path and None cleanup
    - If PDF: converts first page to temp PNG and returns that + cleanup path
    - Else: returns (None, None)
    """
    if not course_outcome_file:
        return None, None

    source = Path(course_outcome_file)
    if not source.exists():
        return None, None

    ext = source.suffix.lower()
    if ext in {".png", ".jpg", ".jpeg", ".bmp", ".gif"}:
        return str(source), None

    if ext == ".pdf":
        try:
            import fitz  # PyMuPDF

            pdf_doc = fitz.open(str(source))
            if pdf_doc.page_count == 0:
                pdf_doc.close()
                return None, None

            strong_start_terms = ["COURSE OUTCOMES", "Course Outcomes"]
            fallback_start_terms = ["CO1", "CO 1"]
            end_terms = ["TEXT BOOKS", "REFERENCE BOOKS", "REFERENCES", "UNIT "]
            mapping_terms = [
                "Levels:",
                "Remembering",
                "Applying",
                "Analyze",
                "Cos / Level",
                "COs / Level",
                "Marks",
            ]

            selected_pix = None

            for page_index in range(min(pdf_doc.page_count, 5)):
                page = pdf_doc.load_page(page_index)
                page_rect = page.rect

                start_rects = []
                for term in strong_start_terms:
                    start_rects.extend(page.search_for(term))

                if not start_rects:
                    for term in fallback_start_terms:
                        start_rects.extend(page.search_for(term))

                if not start_rects:
                    continue

                start_y = min(r.y0 for r in start_rects)

                end_candidates = []
                for term in end_terms:
                    for r in page.search_for(term):
                        if r.y0 > start_y + 15:
                            end_candidates.append(r.y0)

                mapping_rects = []
                for term in mapping_terms:
                    for r in page.search_for(term):
                        if r.y0 > start_y + 10:
                            mapping_rects.append(r)

                if end_candidates:
                    end_y = min(end_candidates) - 6
                else:
                    end_y = page_rect.y1 - 8

                if mapping_rects:
                    mapping_end_y = max(r.y1 for r in mapping_rects) + 20
                    end_y = max(end_y, mapping_end_y)
                    end_y = min(end_y, page_rect.y1 - 6)

                top_y = max(page_rect.y0 + 4, start_y - 2)
                if end_y <= top_y + 20:
                    end_y = min(page_rect.y1 - 8, top_y + page_rect.height * 0.45)

                clip = fitz.Rect(page_rect.x0 + 4, top_y, page_rect.x1 - 4, end_y)
                selected_pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), clip=clip, alpha=False)
                break

            if selected_pix is None:
                page = pdf_doc.load_page(0)
                selected_pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)

            temp_png = Path(tempfile.gettempdir()) / f"co_map_{uuid4().hex}.png"
            selected_pix.save(str(temp_png))
            pdf_doc.close()
            return str(temp_png), str(temp_png)
        except Exception:
            return None, None

    return None, None


def _extract_course_outcomes(course_outcome_file: str) -> List[tuple[str, str]]:
    if not course_outcome_file:
        return []

    source = Path(course_outcome_file)
    if not source.exists():
        return []

    ext = source.suffix.lower()
    rows: List[tuple[str, str]] = []

    def add_row(code: str, desc: str) -> None:
        cleaned = (desc or "").strip()
        if not cleaned:
            return
        rows.append((code.strip(), cleaned))

    stop_terms = [
        "TEXT BOOKS",
        "REFERENCE BOOKS",
        "REFERENCES",
        "UNIT ",
        "COURSE OBJECTIVES",
        "LAB COMPONENTS",
        "LABORATORY",
        "ADVANCED DATA STRUCTURES",
    ]

    def is_stop_line(text: str) -> bool:
        upper = text.upper()
        return any(term in upper for term in stop_terms)

    def slice_course_outcome_lines(lines: List[str]) -> List[str]:
        start_idx = None
        for idx, raw in enumerate(lines):
            upper = (raw or "").upper()
            if "COURSE OUTCOMES" in upper or re.match(r"^CO\s*\d+", upper):
                start_idx = idx
                break

        if start_idx is None:
            return lines

        sliced = []
        for raw in lines[start_idx:]:
            if is_stop_line(raw):
                break
            sliced.append(raw)
        return sliced

    def parse_lines(lines: List[str]) -> None:
        lines = slice_course_outcome_lines(lines)
        current_code = None
        current_desc: List[str] = []

        for raw in lines:
            line = (raw or "").strip()
            if not line:
                continue

            if is_stop_line(line):
                if current_code and current_desc:
                    add_row(current_code, " ".join(current_desc))
                return

            match = re.match(r"^(CO\s*\d+)\s*[:\-\)]?\s*(.*)$", line, re.IGNORECASE)
            if match:
                if current_code and current_desc:
                    add_row(current_code, " ".join(current_desc))
                current_code = match.group(1).upper().replace(" ", "")
                current_desc = [match.group(2).strip()] if match.group(2).strip() else []
                continue

            if current_code:
                current_desc.append(line)

        if current_code and current_desc:
            add_row(current_code, " ".join(current_desc))

    try:
        if ext in {".docx", ".doc"}:
            doc = Document(str(source))
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if not cells or not cells[0]:
                        continue
                    if re.match(r"^CO\s*\d+$", cells[0], re.IGNORECASE):
                        add_row(cells[0].upper().replace(" ", ""), " ".join(cells[1:]))
            if rows:
                return rows

            para_lines = [p.text for p in doc.paragraphs if p.text.strip()]
            parse_lines(para_lines)
            return rows

        if ext == ".pdf":
            try:
                import fitz  # PyMuPDF
            except Exception:
                return []

            pdf_doc = fitz.open(str(source))
            text_lines: List[str] = []
            for page_index in range(min(pdf_doc.page_count, 5)):
                page = pdf_doc.load_page(page_index)
                text = page.get_text("text")
                text_lines.extend(text.splitlines())
            pdf_doc.close()

            parse_lines(text_lines)
            return rows
    except Exception:
        return []

    return rows


def load_blueprint(blueprint_path: str) -> Dict[str, Any]:
    """Load blueprint from JSON file or parse DOCX file"""
    # Check if it's a JSON file
    if blueprint_path.lower().endswith('.json'):
        try:
            with open(blueprint_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            with open(blueprint_path, 'r', encoding='latin-1') as f:
                return json.load(f)
    
    # If it's a DOCX/DOC file, return a default blueprint structure
    elif blueprint_path.lower().endswith(('.docx', '.doc')):
        # Return a default blueprint structure for DOCX files
        return {
            "total_marks": 100,
            "parts": [
                {
                    "name": "Part A",
                    "count": 10,
                    "marks_per_question": 2,
                    "difficulty": "easy",
                    "description": "Multiple Choice Questions"
                },
                {
                    "name": "Part B",
                    "count": 5,
                    "marks_per_question": 5,
                    "difficulty": "medium",
                    "description": "Short Answer Questions"
                },
                {
                    "name": "Part C",
                    "count": 3,
                    "marks_per_question": 10,
                    "difficulty": "hard",
                    "description": "Long Answer Questions"
                }
            ]
        }
    else:
        raise ValueError(f"Unsupported blueprint file format: {blueprint_path}")


def fetch_questions_for_part(
    cursor,
    question_bank_id: int,
    part_name: str,
    count: int,
    difficulty: str = None,
    marks: float = None
) -> List[Dict]:
    """Fetch questions from database matching criteria with fallback logic"""
    
    db_type = get_db_type()
    rand_func = "RANDOM()" if db_type == "sqlite" else "RAND()"
    placeholder = "?" if db_type == "sqlite" else "%s"
    
    print(f"\n  📋 Fetching questions for {part_name}")
    print(f"     Question Bank ID: {question_bank_id}")
    print(f"     Requested count: {count}")
    print(f"     Difficulty: {difficulty}")
    print(f"     Marks: {marks}")
    
    # ✅ Strategy 1: Match by difficulty and marks (NO part name filter)
    query = f"""
        SELECT id, content, part, unit, topic, difficulty, marks 
        FROM questions 
        WHERE question_bank_id = {placeholder}
    """
    params = [question_bank_id]
    
    # Add filters if provided
    if difficulty:
        query += f" AND LOWER(difficulty) = LOWER({placeholder})"
        params.append(difficulty)
    
    if marks:
        query += f" AND ABS(marks - {placeholder}) < 0.5"
        params.append(marks)
    
    query += f" ORDER BY {rand_func} LIMIT {placeholder}"
    params.append(count)
    
    cursor.execute(query, tuple(params))
    result = cursor.fetchall()
    print(f"     Strategy 1 (difficulty + marks): Found {len(result)} questions")
    
    # Strategy 2: If not enough, try with difficulty only (relax marks constraint)
    if len(result) < count and difficulty:
        print(f"     Trying Strategy 2 (difficulty only)...")
        query = f"""
            SELECT id, content, part, unit, topic, difficulty, marks 
            FROM questions 
            WHERE question_bank_id = {placeholder}
            AND LOWER(difficulty) = LOWER({placeholder})
            ORDER BY {rand_func}
            LIMIT {placeholder}
        """
        cursor.execute(query, (question_bank_id, difficulty, count))
        result = cursor.fetchall()
        print(f"     Strategy 2 (difficulty only): Found {len(result)} questions")
    
    # Strategy 3: If still not enough, try marks-based (relax difficulty)
    if len(result) < count and marks:
        print(f"     Trying Strategy 3 (marks-based)...")
        query = f"""
            SELECT id, content, part, unit, topic, difficulty, marks 
            FROM questions 
            WHERE question_bank_id = {placeholder}
            AND ABS(marks - {placeholder}) < 2.0
            ORDER BY {rand_func}
            LIMIT {placeholder}
        """
        cursor.execute(query, (question_bank_id, marks, count))
        result = cursor.fetchall()
        print(f"     Strategy 3 (marks-based): Found {len(result)} questions")
    
    # Strategy 4: Last resort - get any questions from the bank
    if len(result) < count:
        print(f"     Trying Strategy 4 (any questions)...")
        query = f"""
            SELECT id, content, part, unit, topic, difficulty, marks 
            FROM questions 
            WHERE question_bank_id = {placeholder}
            ORDER BY {rand_func}
            LIMIT {placeholder}
        """
        cursor.execute(query, (question_bank_id, count))
        result = cursor.fetchall()
        print(f"     Strategy 4 (any questions): Found {len(result)} questions")
    
    if not result:
        print(f"     ❌ ERROR: No questions found in question bank {question_bank_id}")
    else:
        print(f"     ✅ Successfully fetched {len(result)} questions")
    
    # Convert Row objects to dictionaries
    return [dict(row) for row in result]


def generate_docx_paper(
    title: str,
    subject_name: str,
    exam_type: str,
    exam_date: str,
    total_marks: int,
    duration: str,
    blueprint: Dict,
    questions_by_part: Dict[str, List[Dict]],
    output_path: str,
    course_outcome_file: str = None,
):
    """Generate DOCX question paper"""
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header - SRI SHAKTHI INSTITUTE OF ENGINEERING AND TECHNOLOGY
    header = doc.add_table(rows=1, cols=1, width=Inches(6.5))
    header.style = 'Table Grid'
    cell = header.rows[0].cells[0]
    
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("SRI SHAKTHI INSTITUTE OF ENGINEERING AND TECHNOLOGY")
    r1.bold = True
    r1.font.size = Pt(14)
    
    p2 = cell.add_paragraph("(An Autonomous Institution)")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(10)
    p2.runs[0].bold = True
    
    p3 = cell.add_paragraph("Affiliated to Anna University, Chennai")
    p4 = cell.add_paragraph("Re-Accredited by NAAC with \"A\", Recognized by UGC with Section 2(f) and 12(B)")
    p5 = cell.add_paragraph("NBA Accredited UG Programmes : Agri, BME, BT, CSE, ECE, EEE, MECH, FT and IT")
    p6 = cell.add_paragraph("Coimbatore - 641 062, L & T By Pass, Tamil Nadu, India")
    for p in [p3, p4, p5, p6]:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(8)
    
    doc.add_paragraph()
    
    # Meta / Info Table
    date_str = exam_date if exam_date else datetime.now().strftime("%d.%m.%Y")
    info_table = doc.add_table(rows=1, cols=2)
    info_table.width = Inches(6.5)
    
    # Left side: Date
    p_date = info_table.rows[0].cells[0].paragraphs[0]
    p_date.add_run(f"Date: {date_str}").font.size = Pt(9)
    
    # Right side: Reg No
    p_reg = info_table.rows[0].cells[1].paragraphs[0]
    p_reg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_reg.add_run("Reg No: ").font.size = Pt(9)
    # Using a nested table for boxes if needed, but for docx let's keep it simple with [ ][ ][ ]
    p_reg.add_run("[  ]" * 12).font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Exam title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title_p.add_run(exam_type.upper())
    r_title.bold = True
    r_title.font.size = Pt(12)
    
    # Subject info
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run(f"Subject: {subject_name}")
    r_sub.bold = True
    r_sub.font.size = Pt(11)

    # Time and marks
    tm_table = doc.add_table(rows=1, cols=2)
    tm_table.width = Inches(6.5)
    
    p_time = tm_table.rows[0].cells[0].paragraphs[0]
    p_time.add_run(f"Time: {duration} hours").font.size = Pt(9)
    
    p_marks = tm_table.rows[0].cells[1].paragraphs[0]
    p_marks.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_marks.add_run(f"Maximum: {total_marks} Marks").font.size = Pt(9)
    
    doc.add_paragraph("_" * 80)
    
    # Instructions
    doc.add_paragraph()
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst.add_run("Answer all the Questions").bold = True
    
    doc.add_paragraph("_" * 80)
    doc.add_paragraph()
    
    # Question sections
    question_number = 1
    used_image_ids = set()  # Track images to avoid duplicates in the document
    
    for part in blueprint.get('parts', []):
        # Support both 'name' and 'part_name' keys
        part_name = part.get('part_name') or part.get('name')
        part_questions = questions_by_part.get(part_name, [])
        
        if not part_questions:
            continue
        
        # Part heading
        part_heading = doc.add_paragraph()
        p_r = part_heading.add_run(f"\n{part_name}")
        p_r.bold = True
        effective_count = _effective_answer_count(part, len(part_questions))
        part_heading.add_run(
            f" ({effective_count} × {part['marks_per_question']} = {effective_count * part['marks_per_question']} marks)"
        )
        part_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Part instruction
        instruction = part.get('instructions') or part.get('instruction')
        if instruction:
            inst_para = doc.add_paragraph(instruction)
            inst_para.italic = True
        
        doc.add_paragraph()
        
        # Questions
        for q in part_questions:
            q_para = doc.add_paragraph()
            q_para.paragraph_format.left_indent = Inches(0.4)
            r_num = q_para.add_run(f"{question_number}. ")
            r_num.bold = True
            q_para.add_run(q['content'])
            
            # Try to fetch and insert image for this question
            try:
                image_data = get_image_for_question(q['content'], used_image_ids, trace_label=f"docx_q{question_number}")
                if image_data and image_data.get('image_blob'):
                    # Track this image to avoid duplicates
                    if image_data.get('id'):
                        used_image_ids.add(image_data.get('id'))
                    
                    # Fix extreme brightness before saving
                    fixed_blob = fix_extreme_brightness_image(image_data['image_blob'])
                    # DB images are stored with wrong orientation; rotate while inserting into PDF.
                    fixed_blob = rotate_image_for_pdf_insertion(fixed_blob, clockwise_degrees=90)
                    
                    # Save image to temp file
                    img_temp_path = save_image_blob_to_temp(fixed_blob)
                    if img_temp_path:
                        try:
                            # Add image to document
                            img_para = doc.add_paragraph()
                            img_para.paragraph_format.left_indent = Inches(0.6)
                            run = img_para.add_run()
                            run.add_picture(img_temp_path, width=Inches(4.0))
                            
                            # Add small caption
                            caption = doc.add_paragraph()
                            caption.paragraph_format.left_indent = Inches(0.6)
                            caption_run = caption.add_run(f"[{image_data.get('source_type', 'image')}]")
                            caption_run.font.size = Pt(8)
                            caption_run.italic = True
                            
                            # Mark for cleanup in co_temp_cleanup section
                            logger.info(f"Added image for question {question_number}")
                        except Exception as e:
                            logger.error(f"Error adding image to DOCX: {e}")
                        finally:
                            # Schedule cleanup
                            try:
                                cleanup_temp_image_file(img_temp_path)
                            except:
                                pass
            except Exception as e:
                logger.error(f"Error fetching image for DOCX question {question_number}: {e}")
            
            question_number += 1
            doc.add_paragraph()

    co_rows = _extract_course_outcomes(course_outcome_file)
    co_image_path = None
    co_temp_cleanup = None
    if course_outcome_file and Path(course_outcome_file).exists():
        co_image_path, co_temp_cleanup = _resolve_course_outcome_image(course_outcome_file)

    # Course outcomes section (always show header)
    doc.add_paragraph()
    co_heading = doc.add_paragraph()
    co_heading.add_run("Course Outcomes").bold = True
    if co_rows:
        co_table = doc.add_table(rows=1, cols=2)
        co_table.style = 'Table Grid'
        header_cells = co_table.rows[0].cells
        header_cells[0].text = 'CO'
        header_cells[1].text = 'Course Outcome'
        for code, desc in co_rows:
            row_cells = co_table.add_row().cells
            row_cells[0].text = code
            row_cells[1].text = desc
    elif course_outcome_file and Path(course_outcome_file).exists():
        if co_image_path:
            try:
                doc.add_picture(co_image_path, width=Inches(6.3))
            except Exception:
                doc.add_paragraph(f"Course outcome file uploaded at: {course_outcome_file}")
        else:
            doc.add_paragraph(f"Course outcome file uploaded at: {course_outcome_file}")
    else:
        doc.add_paragraph("Course outcome file not available.")

    blooms_mapping = _build_blooms_mapping(blueprint, questions_by_part)
    doc.add_paragraph()
    blooms_heading = doc.add_paragraph()
    blooms_heading.add_run("Question Blooms Taxonomy Mapping").bold = True

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Level'
    header_cells[1].text = 'Blooms'
    header_cells[2].text = 'Question Nos'
    header_cells[3].text = 'Marks'

    for code in BLOOMS_ORDER:
        row = table.add_row().cells
        row[0].text = code
        row[1].text = blooms_mapping[code]["label"]
        row[2].text = ",".join(str(n) for n in blooms_mapping[code]["questions"])
        row[3].text = f"{blooms_mapping[code]['marks']:.1f}" if blooms_mapping[code]["marks"] else ""

    doc.add_paragraph("_" * 80)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("*** End of Question Paper ***").italic = True
    
    # Save
    doc.save(output_path)

    if co_temp_cleanup:
        try:
            Path(co_temp_cleanup).unlink(missing_ok=True)
        except Exception:
            pass


def generate_pdf_paper(
    title: str,
    subject_name: str,
    exam_type: str,
    exam_date: str,
    total_marks: int,
    duration: str,
    blueprint: Dict,
    questions_by_part: Dict[str, List[Dict]],
    output_path: str,
    course_outcome_file: str = None,
    subject_code: str | None = None,
):
    """Generate PDF question paper"""

    co_image_path = None
    co_temp_cleanup = None
    if course_outcome_file and Path(course_outcome_file).exists():
        co_image_path, co_temp_cleanup = _resolve_course_outcome_image(course_outcome_file)
    
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Check for logo file
    logo_path = Path(__file__).resolve().parent.parent / "data" / "static" / "logo.png"
    has_logo = logo_path.exists()
    
    # Styles
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CollegeTitle',
        parent=styles['Heading1'],
        fontSize=15,
        textColor=colors.black,
        alignment=1, # Center
        fontName='Helvetica-Bold',
        spaceAfter=2
    )
    
    sub_title_style = ParagraphStyle(
        'SubTitle',
        parent=styles['Normal'],
        fontSize=10,
        alignment=1, # Center
        fontName='Helvetica-Bold',
        spaceAfter=1
    )
    
    address_style = ParagraphStyle(
        'Address',
        parent=styles['Normal'],
        fontSize=8,
        alignment=1, # Center
        fontName='Helvetica',
        spaceAfter=1
    )
    
    exam_title_style = ParagraphStyle(
        'ExamTitle',
        parent=styles['Heading2'],
        fontSize=12,
        alignment=1, # Center
        fontName='Helvetica-Bold',
        spaceAfter=4,
        leading=14
    )

    heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading3'],
        fontSize=11,
        fontName='Helvetica-Bold',
        spaceAfter=4,
        leading=12
    )
    
    normal_style = styles['Normal']
    small_style = ParagraphStyle('Small', parent=styles['Normal'], fontSize=9)
    
    # 🏛️ HEADER TABLE (Logo + College Information + NAAC)
    college_info = [
        [Paragraph("SRI SHAKTHI INSTITUTE OF ENGINEERING AND TECHNOLOGY", title_style)],
        [Paragraph("(An Autonomous Institution)", sub_title_style)],
        [Paragraph("Affiliated to Anna University, Chennai", address_style)],
        [Paragraph("Re-Accredited by NAAC with \"A\", Recognized by UGC with Section 2(f) and 12(B)", address_style)],
        [Paragraph("NBA Accredited UG Programmes : Agri, BME, BT, CSE, ECE, EEE, MECH, FT and IT", address_style)],
        [Paragraph("Coimbatore - 641 062, L & T By Pass, Tamil Nadu, India", address_style)],
    ]
    
    info_col = Table(college_info, colWidths=[5.2*inch])
    info_col.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
        ('TOPPADDING', (0, 0), (-1, -1), 1),
    ]))

    # Main header container
    naac_path = Path(__file__).resolve().parent.parent / "data" / "static" / "naac.png"
    has_naac = naac_path.exists()

    from reportlab.platypus import Image
    row_data = []
    col_widths = []

    if has_logo:
        logo_img = Image(str(logo_path), 0.85*inch, 0.85*inch)
        row_data.append(logo_img)
        col_widths.append(0.9*inch)
    else:
        row_data.append("")
        col_widths.append(0.9*inch)

    row_data.append(info_col)
    col_widths.append(5.2*inch)

    if has_naac:
        naac_img = Image(str(naac_path), 0.85*inch, 0.85*inch)
        row_data.append(naac_img)
        col_widths.append(0.9*inch)
    else:
        row_data.append("")
        col_widths.append(0.9*inch)

    header_table = Table([row_data], colWidths=col_widths)
    header_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOX', (0, 0), (-1, -1), 1.5, colors.black),
        ('LEFTPADDING', (0, 0), (-1, -1), 2),
        ('RIGHTPADDING', (0, 0), (-1, -1), 2),
    ]))
    
    elements.append(header_table)
    elements.append(Spacer(1, 0.1*inch))
    
    # 📝 DATE AND REG NO
    date_str = exam_date if exam_date else datetime.now().strftime("%d.%m.%Y")
    
    # Reg No Boxes implementation
    reg_no_cells = [[" " for _ in range(12)]]
    reg_table = Table(reg_no_cells, colWidths=[0.2*inch]*12, rowHeights=[0.25*inch])
    reg_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    meta_data = [
        [Paragraph(f"<b>Date:</b> {date_str}", small_style), 
         Paragraph("<b>Reg No</b>", small_style), reg_table]
    ]
    meta_table = Table(meta_data, colWidths=[2.5*inch, 1*inch, 3*inch])
    meta_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
    ]))
    
    elements.append(meta_table)
    elements.append(Spacer(1, 0.1*inch))
    
    # 🎓 EXAM DETAILS (match college style header)
    exam_title_text = exam_type.upper() if "EXAMINATION" in exam_type.upper() else f"{exam_type.upper()} EXAMINATION"
    elements.append(Paragraph(exam_title_text, exam_title_style))

    # Optional second line from title (e.g., semester/regulation)
    if title:
        elements.append(Paragraph(title, sub_title_style))

    # Subject code + name line like "21OCE02 – Disaster Preparedness and Management"
    if subject_code:
        elements.append(Paragraph(f"{subject_code} \\u2013 {subject_name}", sub_title_style))
    else:
        elements.append(Paragraph(f"Subject: {subject_name}", sub_title_style))
    
    # Time and Marks line
    time_marks = [
        [Paragraph(f"Time: {duration} hours", small_style), 
         Paragraph(f"Maximum: {total_marks} Marks", small_style)]
    ]
    tm_table = Table(time_marks, colWidths=[3.5*inch, 3.5*inch])
    tm_table.setStyle(TableStyle([
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
    ]))
    elements.append(tm_table)
    elements.append(Spacer(1, 0.1*inch))
    elements.append(Paragraph("_" * 105, normal_style))
    elements.append(Spacer(1, 0.15*inch))
    
    # Instructions
    elements.append(Paragraph("<b>Answer all the Questions</b>", ParagraphStyle('Centered', parent=normal_style, alignment=1)))
    elements.append(Spacer(1, 0.1*inch))
    
    # Questions
    question_number = 1
    used_image_ids = set()  # Track images to avoid duplicates in the paper
    
    for part in blueprint.get('parts', []):
        # Support both 'name' and 'part_name' keys
        part_name = part.get('part_name') or part.get('name')
        part_questions = questions_by_part.get(part_name, [])
        
        if not part_questions:
            continue
        
        # Part heading
        effective_count = _effective_answer_count(part, len(part_questions))
        part_text = (
            f"<b>{part_name}</b> ({effective_count} × {part['marks_per_question']} = "
            f"{effective_count * part['marks_per_question']} marks)"
        )
        elements.append(Paragraph(part_text, sub_title_style))
        
        instruction = part.get('instructions') or part.get('instruction')
        if instruction:
            elements.append(Paragraph(f"<i>{instruction}</i>", normal_style))
        
        elements.append(Spacer(1, 0.1*inch))
        
        # Questions
        for q in part_questions:
            q_text = f"<b>{question_number}.</b> {q['content']}"
            elements.append(Paragraph(q_text, normal_style))
            elements.append(Spacer(1, 0.1*inch))
            
            # Try to fetch and insert image for this question
            temp_image_paths = []
            try:
                image_data = get_image_for_question(q['content'], used_image_ids, trace_label=f"pdf_q{question_number}")
                if image_data and image_data.get('image_blob'):
                    # Track this image to avoid duplicates
                    if image_data.get('id'):
                        used_image_ids.add(image_data.get('id'))
                    
                    # Fix extreme brightness before saving
                    fixed_blob = fix_extreme_brightness_image(image_data['image_blob'])
                    # DB images are stored with wrong orientation; rotate while inserting into PDF.
                    fixed_blob = rotate_image_for_pdf_insertion(fixed_blob, clockwise_degrees=90)
                    logger.info(f"PDF question {question_number}: rotation check completed (clockwise=90)")
                    
                    # Save image to temp file
                    img_temp_path = save_image_blob_to_temp(fixed_blob)
                    if img_temp_path:
                        try:
                            # Add image to PDF
                            img_reader = ImageReader(img_temp_path)
                            img_w, img_h = img_reader.getSize()
                            
                            # Fit image inside a safe bounding box to avoid PDF layout overflow
                            # after rotation (portrait images can become very tall).
                            max_width = 5.5 * inch
                            max_height = 3.2 * inch
                            if img_w > 0 and img_h > 0:
                                width_ratio = max_width / float(img_w)
                                height_ratio = max_height / float(img_h)
                                scale = min(width_ratio, height_ratio)
                                draw_width = float(img_w) * scale
                                draw_height = float(img_h) * scale
                            else:
                                draw_width = max_width
                                draw_height = 2.0 * inch
                            
                            img = RLImage(img_temp_path, width=draw_width, height=draw_height)
                            logger.info(
                                f"PDF question {question_number}: draw size set to {draw_width:.1f}x{draw_height:.1f} points"
                            )
                            elements.append(img)
                            elements.append(Spacer(1, 0.05*inch))
                            
                            # Add caption
                            caption_text = f"<i>[Image: {image_data.get('description', 'Related image')} - {image_data.get('source_type', 'source')}]</i>"
                            elements.append(Paragraph(caption_text, small_style))
                            elements.append(Spacer(1, 0.1*inch))
                            
                            temp_image_paths.append(img_temp_path)
                            logger.info(f"Added image for question {question_number} in PDF")
                        except Exception as e:
                            logger.error(f"Error adding image to PDF: {e}")
            except Exception as e:
                logger.error(f"Error fetching image for PDF question {question_number}: {e}")
            
            elements.append(Spacer(1, 0.15*inch))
            question_number += 1
            
            # Store temp paths for cleanup
            if temp_image_paths:
                if not hasattr(generate_pdf_paper, '_temp_image_paths'):
                    generate_pdf_paper._temp_image_paths = []
                generate_pdf_paper._temp_image_paths.extend(temp_image_paths)
        
        elements.append(Spacer(1, 0.2*inch))

    # Course outcomes section (always show header)
    co_rows = _extract_course_outcomes(course_outcome_file)
    elements.append(Spacer(1, 0.2*inch))
    elements.append(Paragraph("<b>Course Outcomes</b>", heading_style))
    if co_rows:
        co_table_data = [["CO", "Course Outcome"]]
        for code, desc in co_rows:
            co_table_data.append([code, Paragraph(desc, normal_style)])
        co_table = Table(co_table_data, colWidths=[0.6*inch, 5.9*inch])
        co_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ]))
        elements.append(co_table)
    elif course_outcome_file and Path(course_outcome_file).exists():
        if co_image_path:
            try:
                img_reader = ImageReader(co_image_path)
                img_w, img_h = img_reader.getSize()
                max_w = 6.5 * inch
                ratio = max_w / float(img_w) if img_w else 1
                img = RLImage(co_image_path, width=max_w, height=float(img_h) * ratio)
                elements.append(img)
            except Exception:
                elements.append(Paragraph(f"Course outcome file uploaded at: {course_outcome_file}", normal_style))
        else:
            elements.append(Paragraph(f"Course outcome file uploaded at: {course_outcome_file}", normal_style))
    else:
        elements.append(Paragraph("Course outcome file not available.", normal_style))

    blooms_mapping = _build_blooms_mapping(blueprint, questions_by_part)
    elements.append(Spacer(1, 0.2*inch))
    elements.append(Paragraph("<b>Question Blooms Taxonomy Mapping</b>", heading_style))

    mapping_data = [["Level", "Blooms", "Question Nos", "Marks"]]
    for code in BLOOMS_ORDER:
        mapping_data.append([
            code,
            blooms_mapping[code]["label"],
            ",".join(str(n) for n in blooms_mapping[code]["questions"]),
            f"{blooms_mapping[code]['marks']:.1f}" if blooms_mapping[code]["marks"] else "",
        ])

    mapping_table = Table(mapping_data, colWidths=[0.8*inch, 1.4*inch, 3.5*inch, 0.8*inch])
    mapping_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ('ALIGN', (3, 0), (3, -1), 'CENTER'),
    ]))
    elements.append(mapping_table)

    # Footer
    elements.append(Paragraph("_" * 100, normal_style))
    elements.append(Spacer(1, 0.1*inch))
    footer_style = ParagraphStyle('Footer', parent=normal_style, alignment=1)
    elements.append(Paragraph("<i>*** End of Question Paper ***</i>", footer_style))
    
    # Build PDF
    try:
        doc.build(elements)
    except Exception as e:
        logger.error(f"Failed while building PDF document: {e}", exc_info=True)
        raise
    finally:
        if co_temp_cleanup:
            try:
                Path(co_temp_cleanup).unlink(missing_ok=True)
            except Exception:
                pass
        
        # Clean up temporary image files
        if hasattr(generate_pdf_paper, '_temp_image_paths'):
            for temp_path in generate_pdf_paper._temp_image_paths:
                try:
                    cleanup_temp_image_file(temp_path)
                except Exception as e:
                    logger.warning(f"Error cleaning up temp image: {e}")
            generate_pdf_paper._temp_image_paths = []


def generate_question_paper(
    cursor,
    title: str,
    subject_id: int,
    subject_name: str,
    question_bank_id: int,
    blueprint_path: str,
    exam_type: str,
    exam_date: str,
    duration: str,
    file_format: str,
    output_path: str
) -> tuple[str, Dict]:
    """
    Main function to generate question paper
    
    Returns: (path to generated file, questions fetched)
    """
    
    # Load blueprint
    blueprint = load_blueprint(blueprint_path)
    
    # ✅ Add debug logging
    print(f"\n{'='*60}")
    print(f"🎯 GENERATING QUESTION PAPER")
    print(f"{'='*60}")
    print(f"📄 Title: {title}")
    print(f"📚 Subject: {subject_name} (ID: {subject_id})")
    print(f"🏦 Question Bank ID: {question_bank_id}")
    print(f"📋 Blueprint: {blueprint.get('name', 'Unnamed')}")
    print(f"📊 Parts in blueprint: {len(blueprint.get('parts', []))}")
    print(f"📝 Format: {file_format.upper()}")
    print(f"{'='*60}\n")
    
    # Fetch questions for each part
    questions_by_part = {}
    
    for part in blueprint.get('parts', []):
        # Support both 'name' and 'part_name' keys
        part_name = part.get('part_name') or part.get('name')
        
        # Support both 'count' and 'num_questions' keys
        count = part.get('num_questions') or part.get('count')
        
        difficulty = part.get('difficulty')
        marks = part.get('marks_per_question')
        
        print(f"🔍 Processing: {part_name}")
        print(f"   Requested: {count} questions @ {marks} marks each")
        print(f"   Difficulty: {difficulty or 'Any'}")
        
        questions = fetch_questions_for_part(
            cursor,
            question_bank_id,
            part_name,
            count,
            difficulty,
            marks
        )
        
        if len(questions) < count:
            print(f"   ⚠️  WARNING: Could only fetch {len(questions)}/{count} questions!")
        else:
            print(f"   ✅ Successfully fetched {len(questions)} questions")
        
        questions_by_part[part_name] = questions
    
    # ✅ Check if we have any questions at all
    total_questions = sum(len(qs) for qs in questions_by_part.values())
    print(f"\n{'='*60}")
    print(f"📊 SUMMARY")
    print(f"{'='*60}")
    print(f"Total questions fetched: {total_questions}")
    
    if total_questions == 0:
        raise Exception("❌ No questions found in the question bank! Please add questions first.")
    
    # Calculate total marks based on effective answerable questions (e.g., "Answer any 2")
    total_marks = 0
    for part in blueprint.get('parts', []):
        part_name = part.get('part_name') or part.get('name')
        part_questions = questions_by_part.get(part_name, [])
        effective_count = _effective_answer_count(part, len(part_questions))
        total_marks += effective_count * float(part.get('marks_per_question') or 0)
    
    print(f"Total marks: {total_marks}")
    print(f"Output file: {output_path}")
    print(f"{'='*60}\n")
    
    # Load course outcome file path for this subject
    course_outcome_file = None
    placeholder = "?" if get_db_type() == "sqlite" else "%s"
    try:
        cursor.execute(f"SELECT course_outcome_file FROM subjects WHERE id = {placeholder}", (subject_id,))
        subject_row = cursor.fetchone()
        if subject_row:
            row = dict(subject_row) if not isinstance(subject_row, dict) else subject_row
            course_outcome_file = row.get("course_outcome_file")
    except Exception:
        course_outcome_file = None

    # Generate paper based on format
    if file_format.lower() == 'pdf':
        print("📄 Generating PDF...")
        generate_pdf_paper(
            title, subject_name, exam_type, exam_date or 'TBD',
            total_marks, duration, blueprint, questions_by_part, output_path, course_outcome_file
        )
    elif file_format.lower() == 'docx':
        print("📄 Generating DOCX...")
        generate_docx_paper(
            title, subject_name, exam_type, exam_date or 'TBD',
            total_marks, duration, blueprint, questions_by_part, output_path, course_outcome_file
        )
    else:
        raise ValueError(f"Unsupported file format: {file_format}")
    
    print(f"\n✅ Question paper generated successfully!")
    print(f"📁 Location: {output_path}")
    print(f"{'='*60}\n")
    
    return output_path, questions_by_part